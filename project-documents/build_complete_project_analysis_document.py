from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent

COLORS = {
    "navy": "0B2545",
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "muted": "64748B",
    "fill": "F2F4F7",
    "soft": "F8FAFC",
}


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(9.5)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, COLORS["fill"])
        cell_text(cell, header, True, COLORS["navy"])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph()


def setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, COLORS["blue"], 16, 8),
        ("Heading 2", 13, COLORS["blue"], 12, 6),
        ("Heading 3", 12, COLORS["dark_blue"], 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    return doc


def title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Aivan HRMS Complete Project Analysis")
    r.font.name = "Calibri"
    r.font.size = Pt(22)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(COLORS["navy"])

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("Project details, working logic, workflow, SDLC, software requirement diagrams, and data-flow diagrams for the HRMS portal.")
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(COLORS["muted"])


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def callout(doc: Document, label: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    shade(cell, COLORS["soft"])
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(COLORS["dark_blue"])
    r = p.add_run(body)
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    doc.add_paragraph()


def code_block(doc: Document, heading: str, code: str) -> None:
    doc.add_heading(heading, level=2)
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    shade(cell, "F8FAFC")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(code)
    r.font.name = "Consolas"
    r.font.size = Pt(8)
    doc.add_paragraph()


def build_doc() -> Path:
    doc = setup_doc()
    title(doc)

    doc.add_heading("1. Project Details", level=1)
    callout(
        doc,
        "Project purpose",
        "Aivan HRMS is a Human Resource Management System used by admin, HR, and employee users to manage accounts, profiles, attendance, dashboards, notifications, login activity, and password recovery.",
    )
    add_table(
        doc,
        ["Area", "Details"],
        [
            ["Frontend", "Angular and TypeScript application with role-based routes, dashboards, auth pages, employee/HR screens, profile pages, attendance screens, services, guards, and UI components."],
            ["Backend", "FastAPI Python API with routes, services, schemas, SQLAlchemy models, repositories, seed data, WebSocket manager, scheduler service, and Alembic migrations."],
            ["Database", "Stores users, roles, HR users, employees, attendance records, daily summaries, notifications, login activity, and approval/time-off related models."],
            ["Main users", "Admin, HR, Employee, and System processes."],
            ["Communication", "Angular calls FastAPI REST endpoints and connects to a WebSocket endpoint for real-time messages."],
        ],
        [1.55, 4.9],
    )

    doc.add_heading("2. Working Logic", level=1)
    add_table(
        doc,
        ["Module", "Working Logic"],
        [
            ["Authentication", "User submits email/password. Backend verifies password hash, determines role and active dashboard, creates JWT token, logs login activity, and frontend stores the session."],
            ["Role routing", "Admin lands on master dashboard, employee lands on employee dashboard, and HR can choose HR dashboard or employee dashboard."],
            ["Employee management", "Admin/HR can add, list, view, update, and view credentials for employees. Creating an employee also creates a linked login user."],
            ["HR management", "Admin creates HR users. Backend creates user account, HR profile, and shadow employee record so HR can also use employee workflows."],
            ["Attendance", "Employee has one attendance record per day. Punch-in marks Working; punch-out calculates work time, break, overtime, late minutes, early exit, and final status."],
            ["Dashboard", "Backend aggregates workforce counts, attendance status counts, recent records, gender/work-mode breakdowns, upcoming birthdays, and analytics."],
            ["Profile", "User profile is read from Employee/User records. Profile updates sync employee details, user display name, and profile image."],
            ["Notifications", "Backend stores notification records and sends real-time WebSocket messages to the logged-in user."],
            ["Password reset", "Backend creates reset token using JWT and current password hash, sends reset link, then verifies token before saving new password."],
        ],
        [1.55, 4.9],
    )

    doc.add_heading("3. Project Workflow", level=1)
    numbers(
        doc,
        [
            "User opens Angular frontend and reaches login screen.",
            "Frontend AuthService sends login request to FastAPI backend.",
            "Backend validates user, creates JWT session, logs login activity, and returns user details.",
            "Frontend stores token and user data, then route guards allow the correct dashboard.",
            "User performs role-specific actions such as creating HR users, managing employees, punching attendance, viewing dashboards, or updating profile.",
            "Frontend services call backend APIs with Authorization bearer token.",
            "Backend routes validate role, call service-layer business logic, and read/write data through SQLAlchemy.",
            "Backend returns JSON data to frontend. Frontend maps it into tables, cards, badges, charts, forms, and detail screens.",
            "When notification events occur, backend stores notification data and pushes real-time WebSocket messages.",
        ],
    )

    doc.add_heading("4. SDLC", level=1)
    doc.add_paragraph("Recommended SDLC model: Agile Iterative Development.")
    add_table(
        doc,
        ["Phase", "HRMS Activities", "Output"],
        [
            ["Requirement Analysis", "Collect admin, HR, and employee requirements for auth, dashboard, employee, HR, attendance, profile, notification, and audit modules.", "SRS, role matrix, backlog"],
            ["Planning", "Prioritize modules, define sprint scope, assign frontend/backend work, decide API and database changes.", "Sprint plan, release plan"],
            ["Design", "Design Angular routes/UI, FastAPI routes/services, DB schema, auth flow, DFDs, UI/UX, and API contracts.", "Architecture, diagrams, wireframes"],
            ["Development", "Implement frontend components/services and backend routes/services/models/migrations.", "Working modules"],
            ["Testing", "Test login, role access, employee creation, attendance rules, dashboard counts, notifications, and APIs.", "Test results, bug fixes"],
            ["Deployment", "Build Angular app, run backend with Uvicorn, configure DB, CORS, SMTP, env variables, and hosting.", "Live system"],
            ["Maintenance", "Fix bugs, optimize, add reports/time-off/payroll, improve security, and refine UX.", "Updated releases"],
        ],
        [1.25, 3.45, 1.75],
    )

    doc.add_heading("5. Software Requirement Diagram", level=1)
    code_block(
        doc,
        "Text Requirement Diagram",
        """Aivan HRMS Requirements
├─ Functional Requirements
│  ├─ Authentication and Password Reset
│  ├─ Role-Based Dashboards
│  ├─ Employee Management
│  ├─ HR User Management
│  ├─ Attendance Punch-In / Punch-Out
│  ├─ Profile Management
│  ├─ Notifications
│  └─ Login Activity Audit
├─ Non-Functional Requirements
│  ├─ Security
│  ├─ Usability
│  ├─ Maintainability
│  ├─ Performance
│  └─ Scalability
├─ Data Requirements
│  ├─ Users and Roles
│  ├─ Employee and HR Profiles
│  ├─ Attendance and Daily Summary
│  ├─ Notifications
│  └─ Login Activity
└─ Integration Requirements
   ├─ Angular to FastAPI REST APIs
   ├─ JWT Authentication
   ├─ WebSocket Notifications
   ├─ SMTP Email
   └─ SQLAlchemy / Alembic Database Layer""",
    )

    doc.add_heading("6. DFD Level 0 - Context Diagram", level=1)
    doc.add_paragraph("Level 0 shows the whole HRMS as one process and its interaction with external users/systems.")
    add_table(
        doc,
        ["External Entity", "Input to HRMS", "Output from HRMS"],
        [
            ["Admin", "Login, HR user data, employee actions, dashboard requests", "Master dashboard, HR list, employee list, attendance reports"],
            ["HR", "Login, employee details, attendance queries, profile updates", "HR dashboard, employee records, attendance analytics"],
            ["Employee", "Login, punch-in/out, profile updates, password changes", "Employee dashboard, attendance state, profile data, notifications"],
            ["SMTP Mail Service", "Reset email content", "Email delivery status"],
            ["Location Service", "IP/location lookup request", "Approximate location data"],
        ],
        [1.45, 2.55, 2.45],
    )
    code_block(
        doc,
        "DFD Level 0 Mermaid",
        """flowchart LR
    Admin[Admin] --> HRMS((Aivan HRMS System))
    HR[HR User] --> HRMS
    Emp[Employee] --> HRMS
    HRMS --> Admin
    HRMS --> HR
    HRMS --> Emp
    HRMS --> SMTP[SMTP Mail Service]
    SMTP --> HRMS
    HRMS --> LOC[Location Service]
    LOC --> HRMS
    HRMS <--> DB[(HRMS Database)]""",
    )

    doc.add_heading("7. DFD Level 1 - Main System Processes", level=1)
    add_table(
        doc,
        ["Process", "Input", "Data Store Used", "Output"],
        [
            ["1.0 Authentication", "Email, password, dashboard choice, reset request", "Users, Roles, LoginActivity", "JWT session, login result, reset link, audit log"],
            ["2.0 User and Workforce Management", "HR/employee details, update requests", "Users, Employees, HR Users", "Created/updated workforce records and credentials"],
            ["3.0 Attendance Management", "Punch-in/out, work mode, location, image, schedule", "Attendance, DailySummary, Employees", "Today state, history, final status, calculated metrics"],
            ["4.0 Dashboard and Reporting", "Dashboard requests, filters", "Employees, HR Users, Attendance, Users", "Counts, analytics, recent timesheets, reports"],
            ["5.0 Profile Management", "Profile view/update data", "Users, Employees", "Profile details and update confirmation"],
            ["6.0 Notifications", "Event messages, mark-read request", "Notifications, WebSocket connections", "Unread count, notification list, real-time push"],
        ],
        [1.65, 1.65, 1.6, 1.55],
    )
    code_block(
        doc,
        "DFD Level 1 Mermaid",
        """flowchart TD
    Admin[Admin] --> P1[1.0 Authentication]
    HR[HR User] --> P1
    Emp[Employee] --> P1
    P1 <--> D1[(Users / Roles)]
    P1 --> D6[(Login Activity)]

    Admin --> P2[2.0 User and Workforce Management]
    HR --> P2
    P2 <--> D2[(Employees)]
    P2 <--> D3[(HR Users)]
    P2 <--> D1

    Emp --> P3[3.0 Attendance Management]
    HR --> P3
    Admin --> P3
    P3 <--> D4[(Attendance)]
    P3 --> D5[(Daily Summary)]

    Admin --> P4[4.0 Dashboard and Reporting]
    HR --> P4
    Emp --> P4
    P4 --> Admin
    P4 --> HR
    P4 --> Emp
    P4 <--> D2
    P4 <--> D3
    P4 <--> D4

    Admin --> P5[5.0 Profile Management]
    HR --> P5
    Emp --> P5
    P5 <--> D1
    P5 <--> D2

    P3 --> P6[6.0 Notifications]
    P6 <--> D7[(Notifications)]
    P6 --> WS[WebSocket Push]
    WS --> Emp
    WS --> HR
    WS --> Admin""",
    )

    doc.add_heading("8. DFD Level 2 - Authentication Process", level=1)
    numbers(
        doc,
        [
            "User submits login credentials.",
            "System validates email and password against Users table.",
            "System checks role and dashboard access.",
            "For HR without selected dashboard, system returns dashboard selection requirement.",
            "For successful login, system creates JWT token.",
            "System logs success/failure in LoginActivity.",
            "Frontend stores token and session user.",
        ],
    )
    code_block(
        doc,
        "DFD Level 2 Authentication Mermaid",
        """flowchart TD
    U[User] --> A1[1.1 Submit Credentials]
    A1 --> A2[1.2 Validate User]
    A2 <--> D1[(Users)]
    A2 --> A3[1.3 Check Role and Dashboard]
    A3 <--> D2[(Roles)]
    A3 -->|HR needs choice| A4[1.4 Return Dashboard Selection]
    A3 -->|Valid login| A5[1.5 Generate JWT]
    A5 --> A6[1.6 Save Login Activity]
    A6 --> D3[(Login Activity)]
    A5 --> S[Session Response to Frontend]""",
    )

    doc.add_heading("9. DFD Level 2 - Attendance Process", level=1)
    numbers(
        doc,
        [
            "Employee requests today's attendance state.",
            "System creates pre-punch record if no record exists.",
            "Employee submits punch-in with work mode/location/image.",
            "System validates duplicate punch rules and stores punch-in.",
            "Employee submits punch-out.",
            "System calculates working minutes, lunch break, overtime, late/early minutes, and status.",
            "System updates daily summary and creates notification.",
            "Frontend refreshes attendance state and displays updated values.",
        ],
    )
    code_block(
        doc,
        "DFD Level 2 Attendance Mermaid",
        """flowchart TD
    Emp[Employee] --> T1[3.1 Get Today State]
    T1 <--> D1[(Attendance)]
    T1 -->|No record| T2[3.2 Create Pre-Punch Record]
    Emp --> T3[3.3 Punch In]
    T3 --> T4[3.4 Validate Single Punch Rule]
    T4 --> D1
    Emp --> T5[3.5 Punch Out]
    T5 --> T6[3.6 Calculate Metrics and Status]
    T6 --> D1
    T6 --> D2[(Daily Summary)]
    T6 --> T7[3.7 Create Notification]
    T7 --> D3[(Notifications)]
    T7 --> WS[WebSocket Message]
    WS --> Emp""",
    )

    doc.add_heading("10. DFD Level 2 - Employee/HR Management Process", level=1)
    code_block(
        doc,
        "DFD Level 2 Workforce Management Mermaid",
        """flowchart TD
    Admin[Admin/HR] --> M1[2.1 Submit Employee or HR Data]
    M1 --> M2[2.2 Validate Role Permission]
    M2 --> M3[2.3 Check Duplicate Email]
    M3 <--> D1[(Users)]
    M3 --> M4[2.4 Create Login User]
    M4 --> D1
    M4 --> M5[2.5 Create Employee Profile]
    M5 --> D2[(Employees)]
    M4 -->|If HR| M6[2.6 Create HR Profile]
    M6 --> D3[(HR Users)]
    M6 -->|If HR| M7[2.7 Create Shadow Employee]
    M7 --> D2
    M5 --> R[Created / Updated Record Response]""",
    )

    doc.add_heading("11. DFD Data Stores", level=1)
    add_table(
        doc,
        ["Data Store", "Stored Data"],
        [
            ["D1 Users/Roles", "Login email, password hash, display name, role, status, active dashboard, profile image."],
            ["D2 Employees", "Employee code, user link, personal details, contact details, department, designation, shift, work location, status."],
            ["D3 HR Users", "HR code, full name, email, phone, department, designation, status."],
            ["D4 Attendance", "Date, punch-in/out, status, work mode, location, address, image, schedule, task, working minutes, break, overtime."],
            ["D5 Daily Summary", "Daily totals, overtime, late minutes, early leave minutes."],
            ["D6 Notifications", "Type, title, message, reference id, read/unread state, timestamps."],
            ["D7 Login Activity", "User, IP address, user agent, login status, timestamp."],
        ],
        [1.75, 4.7],
    )

    doc.add_heading("12. Final Summary", level=1)
    bullets(
        doc,
        [
            "The HRMS follows a modular Angular + FastAPI architecture.",
            "The main business workflows are authentication, employee/HR management, attendance, dashboards, profile, notifications, and audit tracking.",
            "Agile iterative SDLC is the best model because modules can be planned, built, tested, and improved sprint by sprint.",
            "The DFDs show how data moves between users, frontend actions, backend processes, data stores, notifications, email, and location services.",
        ],
    )

    path = OUT_DIR / "HRMS_Complete_Project_Analysis_SDLC_Requirements_DFD.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    print(build_doc())
